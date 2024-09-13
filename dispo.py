import pandas as pd
import streamlit as st
from docx import Document

# Función para cargar y agrupar el archivo CSV
def cargar_datos(file):
    df = pd.read_csv(file, sep=';')
    df_grouped = df.groupby('ID').first()
    return df_grouped

# Función para obtener la lista de coordinadores, MD asistenciales o investigadores
def obtener_coordinadores(df_grouped, columna):
    coordinadores = df_grouped[columna].dropna().unique()
    return coordinadores

# Función para filtrar estudios por el coordinador, MD asistencial o investigador
def estudios_por_coordinador(df_grouped, coordinador_seleccionado, columna):
    estudios_filtrados = df_grouped[df_grouped[columna] == coordinador_seleccionado]
    
    tabla_estudios = pd.DataFrame({
        'Acrónimo': estudios_filtrados['Acrónimo Estudio'],
        'Número del Comité': estudios_filtrados['Número IRB'],
        'Fase del Estudio': estudios_filtrados.apply(
            lambda row: f"Estado: {'Activo' if '1. Activo' in str(row['Estado general del estudio']) else 'Inactivo'}, "
                        f"Reclutamiento: {'Si' if '1. Si' in str(row['Reclutamiento activo']) else 'No'}", axis=1),
        'Sujetos Tamizados': estudios_filtrados['Total fallas de tamizaje'],
        'Sujetos Activos': estudios_filtrados['Total de activos']
    }).reset_index(drop=True)
    
    return tabla_estudios

# Función para generar el documento en Word
def generar_documento(nombre_persona, categoria, tabla_estudios):
    doc = Document()
    doc.add_heading(f"Reporte para {categoria}", 0)
    doc.add_paragraph(f"Nombre: {nombre_persona}")
    
    if not tabla_estudios.empty:
        doc.add_paragraph("Tabla de estudios asociados:")
        table = doc.add_table(rows=1, cols=len(tabla_estudios.columns))
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(tabla_estudios.columns):
            hdr_cells[i].text = col_name
        
        for index, row in tabla_estudios.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)
    else:
        doc.add_paragraph("No hay estudios asociados.")
    
    filename = f"Reporte_{nombre_persona.replace(' ', '_')}.docx"
    doc.save(filename)
    st.success(f"Documento guardado como {filename}")
    with open(filename, "rb") as file:
        st.download_button(label="Descargar reporte", data=file, file_name=filename)
st.write(
    """
    <style>
    .fixed-text {
        position: fixed;
        left: 10px;
        bottom: 10px;
        font-size: 12px;
    }
    </style>
    <div class="fixed-text">
        Desarrollado por: Unidad de Inteligencia Artificial
    </div>
    """,
    unsafe_allow_html=True
)
# Inicialización de la app con Streamlit
st.title("Generador de Informes")
st.header("Centro de Investigaciones Clínicas")
st.subheader("Fundación Valle del Lili")

# Subida del archivo CSV
uploaded_file = st.file_uploader("Cargar base de datos CSV", type=["csv"])

if uploaded_file:
    df_grouped = cargar_datos(uploaded_file)

    # Selección de categoría
    categorias = [
        "Seleccionar", "Coordinador Principal", "Coordinador backup principal 1", "Coordinador backup principal 2",
        "Coordinador backup principal 3", "Coordinador backup principal 4", "Coordinador backup principal 5", 
        "MD asistencial 1", "MD asistencial 2", "MD asistencial 3", "MD asistencial 4", "MD asistencial 5", 
        "MD asistencial 6", "MD asistencial 7", "MD asistencial 8", "Investigador Principal", 
        "Co-Investigador 1", "Co-Investigador 2", "Co-Investigador 3", "Co-Investigador 4", "Co-Investigador 5",
        "Co-Investigador 6", "Co-Investigador 7"
    ]

    seleccion_categoria = st.selectbox("Selecciona una categoría", categorias)

    if seleccion_categoria != "Seleccionar":
        # Obtener coordinadores según la categoría seleccionada
        categoria = seleccion_categoria
        coordinadores = obtener_coordinadores(df_grouped, categoria)

        if len(coordinadores) > 0:
            seleccion_persona = st.selectbox("Seleccionar Persona", coordinadores)
            
            # Botón para generar el informe
            if st.button("Generar Informe"):
                tabla_resultante = estudios_por_coordinador(df_grouped, seleccion_persona, categoria)
                generar_documento(seleccion_persona, categoria, tabla_resultante)
        else:
            st.error(f"No hay personal registrado en la categoría {seleccion_categoria}.")
    else:
        st.warning("Por favor selecciona una categoría.")
